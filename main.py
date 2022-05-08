
import os

from docx2pdf import convert

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm

d = "files"

for path in os.listdir(d):
    fullPath = os.path.join(d, path)
    if os.path.isfile(fullPath):
        document = Document(fullPath)
        for para in document.paragraphs:
            TStyle, FStyle = document.styles['Normal'], document.styles['Heading 1']
            for style in (TStyle, FStyle):
                style.font.name = "Arial"
            TStyle.font.size = Pt(48)
            FStyle.font.size = Pt(24)

            for para in document.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            sections = document.sections

            for section in sections:
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(1.5)
                section.right_margin = Cm(1.5)
                section.header_distance = Cm(0.5)
                section.footer_distance = Cm(0.5)

                footer = section.footer
                myFooter = footer.paragraphs[0]
                myFooter.text = f"Wprowadz tekst stopki zgodnie z wytycznymi"

        resultFileNamePath = fullPath.replace("files", "result")
        document.save(resultFileNamePath)

        pdfResultFileNamePath = resultFileNamePath.replace("result", "pdf")

        convert(resultFileNamePath, pdfResultFileNamePath.replace("docx", "pdf"))
